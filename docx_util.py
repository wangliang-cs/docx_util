import os.path

from docx import Document
from docxcompose.composer import Composer

import datetime


# refer to: https://blog.csdn.net/m0_46219714/article/details/127694673
def merge_docs(source_filepath_list, output_filepath):
    page_break_doc = Document()
    page_break_doc.add_page_break()

    assert len(source_filepath_list) > 0
    for idx, source_path in enumerate(source_filepath_list):
        print(f"合并[{idx + 1}/{len(source_filepath_list)}]:  {source_path}")
    output_composer = Composer(Document(source_filepath_list[0]))

    for idx, source_path in enumerate(source_filepath_list):
        if idx == 0:
            continue
        output_composer.append(page_break_doc)
        output_composer.append(Document(source_path))

    output_composer.save(output_filepath)


def load_source_list():
    source_filepath_list = []
    source_dir = ""
    with open("./mergelist.txt", "r", encoding="utf-8") as f:
        for idx, line in enumerate(f.read().splitlines()):
            if idx == 0:
                source_dir = line
            else:
                source_filepath_list.append(f"{source_dir}{line}")
    return source_filepath_list


def generate_output_path():
    time_of_execution = datetime.datetime.now().strftime('%Y年%m月%d日%H时%M分%S秒')
    filepath = f"./{time_of_execution}_合并文件.docx"
    count = 1
    while os.path.exists(filepath):
        filepath = f"./{time_of_execution}_合并文件_{count}.docx"
        count = count + 1
    return filepath


def gen_h0(content: str):
    doc = Document()
    doc.add_heading(content, level=0)
    return doc


def gen_h1(content: str):
    doc = Document()
    doc.add_heading(content, level=1)
    return doc


def gen_h2(content: str):
    doc = Document()
    doc.add_heading(content, level=2)
    return doc


def gen_h3(content: str):
    doc = Document()
    doc.add_heading(content, level=3)
    return doc


def gen_h4(content: str):
    doc = Document()
    doc.add_heading(content, level=4)
    return doc


def gen_pa(content: str):
    doc = Document()
    doc.add_paragraph(content)
    return doc


def gen_doc(content: str):
    doc = Document(root_dir + content)
    return doc


def gen_page_break(content: str):
    doc = Document()
    doc.add_page_break()
    return doc


cmds = {"h0": gen_h0, "h1": gen_h1, "h2": gen_h2, "h3": gen_h3, "h4": gen_h4, "pa": gen_pa, "do": gen_doc,
        "pb": gen_page_break}


def produce_doc(output_filepath):
    output_composer = Composer(Document("./样式模板.docx"))
    with open("./word_template_cut_2_sec.txt", "r", encoding="utf-8") as f:
        for idx, line in enumerate(f.read().splitlines()):
            strs = line.split(':', 1)
            if strs[0] == "":
                continue
            print(line)
            output_composer.append(cmds[strs[0]](strs[1]))
    output_composer.save(output_filepath)


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    # 样式模板结束
    root_dir = "C:/Users/liang/Documents/我的坚果云/0_WiP/20180917_群智大项目/20230302_验收绩效评价文档===============/素材/分项/"
    if not os.path.exists(root_dir):
        root_dir = "D:/sync/wangliang/我的坚果云/0_WiP/20180917_群智大项目/20230302_验收绩效评价文档===============/素材/分项/"
    output_filepath = generate_output_path()
    produce_doc(output_filepath)
    # source_filepath_list = load_source_list()
    # output_filepath = generate_output_path()
    # merge_docs(source_filepath_list, output_filepath)
    # print("===========================")
    # print(f"输出: {output_filepath}")
